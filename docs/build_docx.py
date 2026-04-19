#!/usr/bin/env python3
"""
Convert docs/MOROHUB_HANDOFF.md into a Word (.docx) document.

Intentionally minimal — this is not a general-purpose MD→DOCX converter. It
only handles the subset of Markdown used in MOROHUB_HANDOFF.md:
  - ATX headings (# .. ######)
  - Fenced code blocks (``` ... ```)
  - Bullet lists (- foo)
  - Numbered lists (1. foo)
  - Pipe tables (| col | col |)
  - Blockquotes (> foo)
  - Horizontal rules (---)
  - Bold (**...**) and inline code (`...`)

Run: python3 docs/build_docx.py
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

ROOT = Path(__file__).resolve().parent.parent

# (source.md, output.docx, title, subtitle) — add new entries here to emit
# more DOCX files in one run.
DOCS = [
    (
        ROOT / "docs" / "MOROHUB_HANDOFF.md",
        ROOT / "docs" / "MOROHUB_HANDOFF.docx",
        "MoroHub Handoff",
        "Meeting Preparation and Walk-through Guide",
    ),
    (
        ROOT / "docs" / "DEPLOYMENT_GUIDE.md",
        ROOT / "docs" / "DEPLOYMENT_GUIDE.docx",
        "Deployment & Release Guide",
        "MoroHub Kubernetes Service — Release v0.1.0",
    ),
    (
        ROOT / "docs" / "MOROHUB_EMAIL.md",
        ROOT / "docs" / "MOROHUB_EMAIL.docx",
        "Email to MoroHub Team",
        "Copy/paste ready — Subject + Body",
    ),
]


INLINE_CODE = re.compile(r"`([^`]+)`")
BOLD = re.compile(r"\*\*([^*]+)\*\*")
LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def add_runs(paragraph, text: str, *, base_bold: bool = False) -> None:
    """Split text on inline code / bold / links and add styled runs."""
    # Tokenise by finding matches of inline code, bold, links in order.
    pattern = re.compile(
        r"(`[^`]+`)|(\*\*[^*]+\*\*)|(\[[^\]]+\]\([^)]+\))"
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            plain = text[pos:m.start()]
            r = paragraph.add_run(plain)
            if base_bold:
                r.bold = True
        token = m.group(0)
        if token.startswith("`") and token.endswith("`"):
            r = paragraph.add_run(token[1:-1])
            r.font.name = "Menlo"
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xB0, 0x20, 0x60)
        elif token.startswith("**"):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
        elif token.startswith("["):
            link_m = LINK.match(token)
            if link_m:
                r = paragraph.add_run(link_m.group(1))
                r.font.color.rgb = RGBColor(0x0B, 0x5C, 0xFF)
                r.underline = True
        pos = m.end()
    if pos < len(text):
        tail = text[pos:]
        r = paragraph.add_run(tail)
        if base_bold:
            r.bold = True


def add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # Light grey background via shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)

    r = p.add_run(code.rstrip("\n"))
    r.font.name = "Menlo"
    r.font.size = Pt(9)
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Menlo")
    rFonts.set(qn("w:hAnsi"), "Menlo")
    rFonts.set(qn("w:cs"), "Menlo")


def add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)


def parse_table(lines: list[str]) -> list[list[str]]:
    """Parse GitHub-style pipe table into rows (strings)."""
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    # Drop the separator row (contains dashes only)
    rows = [r for r in rows if not all(re.fullmatch(r":?-+:?", c) for c in r)]
    return rows


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid"
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.text = ""
            para = cell.paragraphs[0]
            add_runs(para, cell_text, base_bold=(ri == 0))


def convert(md_path: Path, docx_path: Path, title_text: str, subtitle_text: str) -> None:
    text = md_path.read_text(encoding="utf-8")
    # Strip the top H1 — we add a formatted title ourselves.
    lines = text.splitlines()

    doc = Document()

    # Global style tweaks
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Title block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run(title_text)
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(0x0B, 0x1F, 0x4A)

    subtitle = doc.add_paragraph()
    r = subtitle.add_run(subtitle_text)
    r.italic = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    add_horizontal_rule(doc)

    i = 0
    skip_first_h1 = True
    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # Fenced code block
        if stripped.startswith("```"):
            j = i + 1
            code_lines = []
            while j < len(lines) and not lines[j].rstrip().startswith("```"):
                code_lines.append(lines[j])
                j += 1
            add_code_block(doc, "\n".join(code_lines))
            i = j + 1
            continue

        # Horizontal rule
        if stripped == "---":
            add_horizontal_rule(doc)
            i += 1
            continue

        # Heading
        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            if level == 1 and skip_first_h1:
                skip_first_h1 = False
                i += 1
                continue
            doc_level = min(level, 4)
            h = doc.add_heading(level=doc_level)
            add_runs(h, text, base_bold=True)
            i += 1
            continue

        # Blockquote
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(6)
            add_runs(p, stripped[2:].strip())
            r = p.runs[0] if p.runs else p.add_run("")
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x50, 0x50, 0x50)
            i += 1
            continue

        # Pipe table
        if stripped.startswith("|") and "|" in stripped[1:]:
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i])
                i += 1
            rows = parse_table(tbl_lines)
            add_table(doc, rows)
            doc.add_paragraph()
            continue

        # Unordered list
        if re.match(r"^\s*[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            content = re.sub(r"^\s*[-*]\s+", "", line)
            add_runs(p, content)
            i += 1
            continue

        # Ordered list
        if re.match(r"^\s*\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            content = re.sub(r"^\s*\d+\.\s+", "", line)
            add_runs(p, content)
            i += 1
            continue

        # Blank line
        if stripped == "":
            i += 1
            continue

        # Regular paragraph — collapse consecutive non-empty lines
        para_lines = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(
            r"^(#{1,6}\s|\s*[-*]\s|\s*\d+\.\s|\|.*\||>\s|```|---$)", lines[i]
        ):
            para_lines.append(lines[i].strip())
            i += 1
        p = doc.add_paragraph()
        add_runs(p, " ".join(para_lines))

    doc.save(docx_path)
    print(f"Wrote {docx_path} ({docx_path.stat().st_size:,} bytes)")


if __name__ == "__main__":
    for src, dst, title, subtitle in DOCS:
        convert(src, dst, title, subtitle)
